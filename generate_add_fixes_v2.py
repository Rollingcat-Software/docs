from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from datetime import datetime
import numpy as np

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_gantt_chart():
    """Create professional Gantt chart images"""

    # Fall Semester Tasks
    fall_tasks = [
        ("F-1: Project Setup", "2025-09-01", "2025-09-30"),
        ("F-2: Database Design", "2025-09-15", "2025-10-31"),
        ("F-3: Identity Core API", "2025-10-01", "2025-11-30"),
        ("F-4: Biometric Processor", "2025-10-01", "2025-11-30"),
        ("F-5: Liveness Detection", "2025-11-01", "2025-12-31"),
        ("F-6: Web Dashboard", "2025-11-01", "2025-12-31"),
        ("F-7: Service Integration", "2025-12-01", "2026-01-25"),
        ("F-8: Mobile App", "2025-12-01", "2026-01-25"),
        ("F-9: Desktop App", "2026-01-01", "2026-01-25"),
        ("F-10: NFC Reader", "2026-01-01", "2026-01-25"),
    ]

    # Spring Semester Tasks
    spring_tasks = [
        ("S-1: RBAC Implementation", "2026-02-01", "2026-02-28"),
        ("S-2: Full Integration", "2026-02-01", "2026-03-31"),
        ("S-3: Mobile - Production", "2026-02-01", "2026-03-31"),
        ("S-4: Desktop - Production", "2026-03-01", "2026-04-30"),
        ("S-5: E2E Testing", "2026-03-01", "2026-04-30"),
        ("S-6: Performance Tuning", "2026-04-01", "2026-05-31"),
        ("S-7: Security Audit", "2026-05-01", "2026-06-15"),
    ]

    def create_chart(tasks, title, filename, color):
        fig, ax = plt.subplots(figsize=(12, 6))

        # Parse dates
        task_names = [t[0] for t in tasks]
        starts = [datetime.strptime(t[1], "%Y-%m-%d") for t in tasks]
        ends = [datetime.strptime(t[2], "%Y-%m-%d") for t in tasks]

        # Find date range
        min_date = min(starts)
        max_date = max(ends)

        # Plot bars
        for i, (name, start, end) in enumerate(zip(task_names, starts, ends)):
            duration = (end - start).days
            start_offset = (start - min_date).days
            ax.barh(i, duration, left=start_offset, height=0.6, color=color, edgecolor='black', linewidth=0.5)

        # Configure axes
        ax.set_yticks(range(len(task_names)))
        ax.set_yticklabels(task_names, fontsize=10)
        ax.invert_yaxis()

        # X-axis: months
        total_days = (max_date - min_date).days
        month_ticks = []
        month_labels = []
        current = min_date.replace(day=1)
        while current <= max_date:
            offset = (current - min_date).days
            month_ticks.append(offset)
            month_labels.append(current.strftime("%b %Y"))
            # Move to next month
            if current.month == 12:
                current = current.replace(year=current.year + 1, month=1)
            else:
                current = current.replace(month=current.month + 1)

        ax.set_xticks(month_ticks)
        ax.set_xticklabels(month_labels, rotation=45, ha='right', fontsize=9)
        ax.set_xlim(-5, total_days + 10)

        # Grid and styling
        ax.grid(axis='x', linestyle='--', alpha=0.7)
        ax.set_axisbelow(True)
        ax.set_title(title, fontsize=14, fontweight='bold', pad=15)
        ax.set_xlabel('Timeline', fontsize=11)

        # Remove top and right spines
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

        plt.tight_layout()
        plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close()
        print(f"Created: {filename}")

    # Generate both charts
    create_chart(fall_tasks, "FIVUCSAS - Fall Semester 2025 Timeline", "gantt_fall.png", "#4472C4")
    create_chart(spring_tasks, "FIVUCSAS - Spring Semester 2026 Timeline", "gantt_spring.png", "#ED7D31")

def create_improved_add_fixes():
    # First, create the Gantt charts
    print("Generating Gantt chart images...")
    create_gantt_chart()

    doc = Document()

    # ---------------------------------------------------------
    # STYLE CONFIGURATION (Aptos Narrow 12)
    # ---------------------------------------------------------
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Aptos Narrow'
    font.size = Pt(12)

    for i in range(1, 4):
        try:
            h_style = doc.styles[f'Heading {i}']
            h_style.font.name = 'Aptos Narrow'
        except:
            pass

    # ==========================================================
    # TITLE
    # ==========================================================
    heading = doc.add_heading('ADD Fixes - Final Version', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('This document contains corrected sections for the FIVUCSAS ADD.')
    doc.add_paragraph('Font: Aptos Narrow 12 | Strictly Compliant with CSE4197 ADD Guide')
    doc.add_paragraph('=' * 60)

    # ==========================================================
    # 0. CHARACTER ENCODING FIXES
    # ==========================================================
    doc.add_heading('0. Character Encoding Fixes', level=1)
    doc.add_paragraph('Replace corrupted Turkish characters with these correct versions:')

    encoding_table = doc.add_table(rows=6, cols=2)
    encoding_table.style = 'Table Grid'
    encoding_data = [
        ['Corrupted Text', 'Correct Text'],
        ['Ahmet Abdullah G�ltekin', 'Ahmet Abdullah Gültekin'],
        ['Ay�e G�ls�m Eren', 'Ayşe Gülsüm Eren'],
        ['Ay�enur Ar�c�', 'Ayşenur Arıcı'],
        ['Assoc. Prof. Dr. Mustafa A�ao�lu', 'Assoc. Prof. Dr. Mustafa Ağaoğlu'],
        ['moir� (Section 2.3, 6.1)', 'moiré']
    ]
    for i, row_data in enumerate(encoding_data):
        row = encoding_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(row.cells[j], 'D9E2F3')

    # ==========================================================
    # 1. FUNCTIONAL REQUIREMENTS (Section 3.1)
    # ==========================================================
    doc.add_heading('1. Replacement for Section 3.1 (Functional Requirements)', level=1)
    doc.add_paragraph('Replace the entire Section 3.1 with the following. This version uses the strict "3.1.X.X" numbering as required by the CSE4197 ADD Guide.')
    doc.add_paragraph()

    doc.add_heading('3.1 Functional Requirements', level=2)
    intro = doc.add_paragraph()
    intro.add_run('The following functional requirements are described at an abstract level to emphasize system behavior and constraints, while implementation-specific details are presented in the system design and architecture sections.')

    # Keep original summary table
    doc.add_paragraph()
    fr_summary = [
        ['ID', 'Functional Requirements', 'Description'],
        ['FR-1', 'Identity & Access Management', 'The system shall authenticate users and manage accounts securely within tenant boundaries.'],
        ['FR-2', 'Biometric Enrollment', 'The system shall enroll facial biometric data after validating quality and liveness.'],
        ['FR-3', 'Biometric Verification', 'The system shall verify user identity using facial biometrics with optional liveness checks.'],
        ['FR-4', 'Multi-Tenant Management', 'The system shall support isolated tenant configuration, quotas, and settings.'],
        ['FR-5', 'Authorization & RBAC', 'The system shall enforce role-based access control for all protected operations.'],
        ['FR-6', 'Auditing & Compliance', 'The system shall record and expose audit logs and verification history securely.']
    ]

    summary_table = doc.add_table(rows=len(fr_summary), cols=3)
    summary_table.style = 'Table Grid'
    for i, row_data in enumerate(fr_summary):
        row = summary_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(row.cells[j], 'D9E2F3')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Table 3: Functional Requirements (FR)').italic = True
    doc.add_paragraph()

    # Detailed FRs with required structure
    frs = [
        {
            "id": "FR-1",
            "title": "Identity and Access Management",
            "desc": "The system shall provide a secure mechanism for authenticating users and managing administrative accounts within strictly isolated tenant boundaries.",
            "inputs": [
                "User credentials (email address, password)",
                "Tenant identifiers",
                "Session-related information (IP address, User-Agent, timestamp)"
            ],
            "proc": [
                "Validate credential formats (RFC 5322 for email, minimum password requirements)",
                "Enforce user uniqueness within each tenant namespace",
                "Verify password against stored BCrypt hash (work factor 12)",
                "Generate signed JWT access token (15-minute expiry) and UUID-based refresh token (7-day expiry)",
                "Store session metadata in Redis cache"
            ],
            "out": [
                "Authentication tokens (Access Token, Refresh Token)",
                "User profile information (ID, email, roles, tenant context)",
                "Session expiration timestamps"
            ],
            "err": [
                "Invalid Credentials: HTTP 401 Unauthorized with generic error message",
                "Duplicate Account: HTTP 409 Conflict during registration",
                "Locked Account: HTTP 423 Locked after 5 consecutive failed attempts",
                "Invalid Token: HTTP 401 Unauthorized for expired or malformed tokens"
            ]
        },
        {
            "id": "FR-2",
            "title": "Biometric Enrollment",
            "desc": "The system shall support the secure enrollment of facial biometric data after verifying both image quality and user liveness.",
            "inputs": [
                "Facial images captured during interactive liveness challenge",
                "User identifier (from authenticated session)",
                "Tenant context"
            ],
            "proc": [
                "Detect facial region using configurable backend (RetinaFace, MTCNN, or MediaPipe)",
                "Assess biometric sample quality (brightness > 0.3, sharpness > 0.5, frontal pose within 15 degrees)",
                "Validate liveness through 'Biometric Puzzle' challenge sequence",
                "Generate standardized biometric embedding using selected model (default: Facenet512)",
                "Store embedding vector in PostgreSQL with pgvector extension"
            ],
            "out": [
                "Unique enrollment identifier (UUID)",
                "Quality confidence score (0.0 - 1.0)",
                "Enrollment status (SUCCESS, PENDING, FAILED)"
            ],
            "err": [
                "No Face Detected: HTTP 400 Bad Request with guidance message",
                "Image Quality Insufficient: HTTP 422 Unprocessable Entity with specific quality metrics",
                "Liveness Validation Failed: HTTP 403 Forbidden with failure reason (e.g., 'Blink not detected')"
            ]
        },
        {
            "id": "FR-3",
            "title": "Biometric Verification",
            "desc": "The system shall verify user identity by comparing live facial biometric samples against previously enrolled biometric templates, with optional liveness validation.",
            "inputs": [
                "Live facial image (base64 encoded or multipart upload)",
                "User or tenant context for template retrieval",
                "Optional: Liveness challenge response"
            ],
            "proc": [
                "Generate biometric embedding from input image using configured model",
                "Retrieve enrolled reference template(s) from vector database",
                "Compute cosine similarity score between input and reference embeddings",
                "Compare similarity score against configurable threshold (default: 0.6)",
                "Log verification attempt with result and metadata"
            ],
            "out": [
                "Verification decision (MATCH / NO_MATCH)",
                "Similarity confidence score (0.0 - 1.0)",
                "Matched identity information (if successful)"
            ],
            "err": [
                "No Enrollment Found: HTTP 404 Not Found",
                "Verification Failed: HTTP 401 Unauthorized when score < threshold",
                "Rate Limit Exceeded: HTTP 429 Too Many Requests"
            ]
        },
        {
            "id": "FR-4",
            "title": "Multi-Tenant Management",
            "desc": "The system shall support isolated configurations for multiple tenants, enabling independent management of quotas, limits, and operational settings.",
            "inputs": [
                "Tenant metadata (organization name, domain, contact email)",
                "Configuration parameters (max users, enrollment quota, API rate limits)",
                "Billing tier selection"
            ],
            "proc": [
                "Validate tenant identifier uniqueness across the platform",
                "Provision tenant-specific configuration with default quotas",
                "Generate tenant API credentials (API Key, Secret)",
                "Apply row-level security policies for data isolation"
            ],
            "out": [
                "Unique tenant identifier (UUID)",
                "Tenant API credentials",
                "Configuration confirmation with applied quotas"
            ],
            "err": [
                "Duplicate Tenant: HTTP 409 Conflict",
                "Quota Exceeded: HTTP 429 Too Many Requests when tenant limits are reached",
                "Invalid Configuration: HTTP 422 Unprocessable Entity for invalid parameters"
            ]
        },
        {
            "id": "FR-5",
            "title": "Authorization & Role Based Access Control",
            "desc": "The system shall enforce role-based access control to regulate access to protected resources based on assigned permissions.",
            "inputs": [
                "Requesting user identity (extracted from JWT)",
                "Target resource identifier (URI path)",
                "Requested operation (HTTP method: GET, POST, PUT, DELETE)"
            ],
            "proc": [
                "Decode and validate JWT token signature and expiration",
                "Extract user roles from token claims",
                "Resolve role-to-permission mappings from RBAC configuration",
                "Evaluate if user permissions include the requested operation on the target resource",
                "Cache permission decisions for performance (TTL: 5 minutes)"
            ],
            "out": [
                "Authorization decision (ALLOW / DENY)",
                "Effective permissions list (for debugging/audit)"
            ],
            "err": [
                "Insufficient Permissions: HTTP 403 Forbidden",
                "Invalid Token: HTTP 401 Unauthorized",
                "Resource Not Found: HTTP 404 Not Found"
            ]
        },
        {
            "id": "FR-6",
            "title": "Auditing & Compliance",
            "desc": "The system shall record and expose immutable audit logs and verification histories to support security monitoring and compliance with applicable data protection regulations (GDPR, KVKK).",
            "inputs": [
                "User action details (actor, action type, target resource)",
                "Request context (IP address, User-Agent, timestamp)",
                "Affected data identifiers"
            ],
            "proc": [
                "Sanitize input data to exclude sensitive credentials and biometric data",
                "Append ISO 8601 timestamp and tenant context",
                "Generate immutable audit record with unique event ID",
                "Write to append-only audit log storage with tenant isolation"
            ],
            "out": [
                "Paginated audit record list with filtering support",
                "Verification history reports (exportable as CSV/JSON)",
                "Compliance summary dashboards"
            ],
            "err": [
                "Access Denied: HTTP 403 Forbidden (only Tenant Admin and above)",
                "Invalid Filter Parameters: HTTP 400 Bad Request",
                "Export Limit Exceeded: HTTP 429 Too Many Requests for large exports"
            ]
        }
    ]

    for idx, fr in enumerate(frs, 1):
        # 3.1.X Heading
        doc.add_heading(f"3.1.{idx} {fr['title']} ({fr['id']})", level=3)

        # 3.1.X.1 Description
        p = doc.add_paragraph()
        p.add_run(f"3.1.{idx}.1 Description").bold = True
        doc.add_paragraph(fr["desc"])

        # 3.1.X.2 Inputs
        p = doc.add_paragraph()
        p.add_run(f"3.1.{idx}.2 Inputs").bold = True
        for inp in fr["inputs"]:
            doc.add_paragraph(inp, style='List Bullet')

        # 3.1.X.3 Processing
        p = doc.add_paragraph()
        p.add_run(f"3.1.{idx}.3 Processing").bold = True
        for i, proc in enumerate(fr["proc"], 1):
            doc.add_paragraph(f"{i}. {proc}")

        # 3.1.X.4 Outputs
        p = doc.add_paragraph()
        p.add_run(f"3.1.{idx}.4 Outputs").bold = True
        for out in fr["out"]:
            doc.add_paragraph(out, style='List Bullet')

        # 3.1.X.5 Error/Data Handling
        p = doc.add_paragraph()
        p.add_run(f"3.1.{idx}.5 Error/Data Handling").bold = True
        for err in fr["err"]:
            doc.add_paragraph(err, style='List Bullet')

        doc.add_paragraph()

    # ==========================================================
    # 2. NON-FUNCTIONAL REQUIREMENTS (Section 3.2)
    # ==========================================================
    doc.add_heading('2. Replacement for Section 3.2 (Non-Functional Requirements)', level=1)
    doc.add_paragraph('Replace Section 3.2 with this version. Scalability is now a separate row, and detailed paragraphs are preserved.')
    doc.add_paragraph()

    doc.add_heading('3.2 Non-Functional Requirements', level=2)

    # NFR Summary Table
    nfr_data = [
        ['Category', 'Requirement', 'Measure / Verification'],
        ['Performance', 'The system shall respond to requests within acceptable latency.', '95th percentile latency < 200ms for authentication; < 500ms for biometric verification.'],
        ['Scalability', 'The system shall handle concurrent users without degradation.', 'Load testing results showing stability up to 100 concurrent users per instance.'],
        ['Reliability', 'The system shall remain available during operational hours.', 'Targeted uptime of 99.5% with an RPO of < 1 hour for backups.'],
        ['Security', 'The system shall protect data against unauthorized access.', 'TLS 1.3 encryption for transmission and AES-256 encrypted storage for embeddings.'],
        ['Usability', 'Users shall complete enrollment with minimal effort.', 'Total enrollment time < 60 seconds including liveness challenges.'],
        ['Maintainability', 'The system shall be modular and easy to extend.', 'Use of Hexagonal Architecture and > 70% unit test coverage.'],
        ['Portability', 'The system shall run consistently across platforms.', 'Docker containerization and KMP shared code verification.']
    ]

    nfr_table = doc.add_table(rows=len(nfr_data), cols=3)
    nfr_table.style = 'Table Grid'
    for i, row_data in enumerate(nfr_data):
        row = nfr_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(row.cells[j], 'D9E2F3')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Table 4: Non-Functional Requirements (NFR)').italic = True
    doc.add_paragraph()

    doc.add_paragraph('The following requirements specify the operational criteria and constraints that define the system\'s quality and performance. These are categorized by their specific attributes to ensure the platform is robust, secure, and maintainable.')
    doc.add_paragraph()

    # NFR Detailed Paragraphs
    nfr_details = [
        ("3.2.1 Performance", "The system shall provide responsive system behavior for authentication and biometric operations in order to support real-time usage scenarios such as physical access control, kiosk authentication, and online identity verification. Performance requirements are critical due to the computational nature of biometric processing and the user-facing interaction model of the system. The platform shall ensure that authentication, enrollment, and verification workflows complete within acceptable time bounds so that biometric verification does not introduce perceptible delays or degrade user experience. Performance compliance shall be evaluated through controlled testing under representative workloads."),
        ("3.2.2 Scalability", "The system shall be designed to scale horizontally in order to support multiple tenants, increasing user populations, and growing biometric workloads. As a Software-as-a-Service platform targeting B2B and B2B2C use cases, the system must accommodate variations in tenant size and request volume without requiring architectural redesign. Scalability requirements ensure that biometric enrollment and verification services can handle concurrent requests across tenants while preserving isolation and performance. Scalability shall be validated through load and stress testing aligned with projected usage scenarios."),
        ("3.2.3 Reliability", "The system shall operate reliably during designated operational periods, ensuring consistent availability of identity verification services. Reliability is a critical requirement due to the system's role in security-sensitive workflows, where service interruptions may block user access or disrupt operations. The system shall tolerate partial failures and recover gracefully without cascading effects across components or tenants. Reliability targets shall be assessed through availability monitoring, fault injection, and controlled failure scenarios."),
        ("3.2.4 Security", "The system shall enforce strict security controls to protect identity information and biometric data throughout their lifecycle. Security requirements address confidentiality, integrity, and access control for sensitive data, particularly given the irreversible nature of biometric identifiers. The system shall prevent unauthorized access, ensure strong isolation between tenants, and minimize exposure of biometric data by operating on derived representations rather than raw biometric artifacts. Compliance with security requirements shall be verified through security testing, auditability, and adherence to best practices in secure system design."),
        ("3.2.5 Usability", "The system shall provide an intuitive and efficient user experience across all supported client applications, including web, mobile, and desktop environments. Usability requirements are especially important for biometric interactions, where unclear feedback or complex flows can lead to user frustration or enrollment failure. The system shall guide users through enrollment and verification processes with clear instructions and meaningful feedback, enabling successful completion with minimal effort. Usability effectiveness shall be evaluated based on task completion success and consistency of interaction flows."),
        ("3.2.6 Maintainability", "The system shall be maintainable and extensible to support future enhancements, algorithm upgrades, and evolving security requirements. Given the rapid advancement of biometric technologies and identity verification standards, the system must allow components to be modified or replaced without disrupting core functionality. Maintainability requirements emphasize modular structure, clear separation of concerns, and ease of testing, enabling the system to evolve throughout its lifecycle. Compliance with maintainability goals shall be assessed through architectural conformance and codebase organization."),
        ("3.2.7 Portability", "The system shall operate consistently across supported deployment environments and client platforms. Portability requirements ensure that the system can be deployed in different infrastructure settings and accessed through multiple operating systems without behavioral inconsistencies. This is essential for a SaaS platform intended for integration into diverse organizational environments. Portability shall be validated through deployment and execution testing across supported platforms and environments.")
    ]

    for title, content in nfr_details:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        doc.add_paragraph(content)

    # ==========================================================
    # 3. TASK LOG (Section 6.2) - With Full Format
    # ==========================================================
    doc.add_heading('3. Replacement for Section 6.2 (Task Log)', level=1)
    p = doc.add_paragraph()
    p.add_run('IMPORTANT: ').bold = True
    p.add_run('Please verify and update the meeting hours below with your actual meeting durations.')
    doc.add_paragraph()

    doc.add_heading('6.2 Task Log', level=2)

    meetings = [
        {
            "num": 1, "date": "12.09.2025", "location": "Online (Microsoft Teams)",
            "period": "[UPDATE HOURS]", "attendees": "All team members, Advisor",
            "objectives": "Project topic selection and team formation",
            "decisions": [
                "Final Engineering Project topic (FIVUCSAS) was selected and approved.",
                "Team roles and responsibilities were assigned.",
                "Project scope and core scenarios were defined.",
                "Initial literature sources on face recognition and biometric systems were identified."
            ]
        },
        {
            "num": 2, "date": "26.09.2025", "location": "Online (Microsoft Teams)",
            "period": "[UPDATE HOURS]", "attendees": "All team members, Advisor",
            "objectives": "Architecture design and technology selection",
            "decisions": [
                "System architecture and microservice structure were designed.",
                "Backend, frontend, and database technologies were selected.",
                "Face recognition and liveness detection literature was reviewed.",
                "Commercial biometric solutions were analyzed for comparison."
            ]
        },
        {
            "num": 3, "date": "10.10.2025", "location": "Online (Microsoft Teams)",
            "period": "[UPDATE HOURS]", "attendees": "All team members, Advisor",
            "objectives": "Database design and multi-tenancy strategy",
            "decisions": [
                "Database schema and multi-tenancy strategy were finalized.",
                "Biometric data storage using face embeddings was defined.",
                "Vector similarity search techniques were evaluated.",
                "Literature on biometric template protection was reviewed."
            ]
        },
        {
            "num": 4, "date": "24.10.2025", "location": "Online (Microsoft Teams)",
            "period": "[UPDATE HOURS]", "attendees": "All team members, Advisor",
            "objectives": "API design and security architecture review",
            "decisions": [
                "Identity Core API design and security architecture were reviewed.",
                "Authentication and authorization mechanisms were finalized.",
                "Secure biometric authentication studies were analyzed.",
                "API versioning and documentation strategy were defined."
            ]
        },
        {
            "num": 5, "date": "07.11.2025", "location": "Online (Microsoft Teams)",
            "period": "[UPDATE HOURS]", "attendees": "All team members, Advisor",
            "objectives": "Biometric processing prototype implementation",
            "decisions": [
                "DeepFace-based biometric processing prototype was implemented.",
                "Biometric enrollment and verification workflows were defined.",
                "Image quality assessment criteria were established.",
                "Performance optimization strategies were discussed."
            ]
        },
        {
            "num": 6, "date": "14.11.2025", "location": "Marmara University, Engineering Faculty",
            "period": "[UPDATE HOURS]", "attendees": "All team members, Advisor",
            "objectives": "Mid-semester demo and progress evaluation",
            "decisions": [
                "Mid-semester system demo was conducted.",
                "Project progress was evaluated against the timeline.",
                "Integration and performance risks were identified.",
                "Accuracy metrics evaluation approach was defined."
            ]
        },
        {
            "num": 7, "date": "28.11.2025", "location": "Online (Microsoft Teams)",
            "period": "[UPDATE HOURS]", "attendees": "All team members, Advisor",
            "objectives": "Liveness detection algorithm evaluation",
            "decisions": [
                "Biometric Puzzle liveness detection algorithm evaluated.",
                "Active and passive liveness detection methods were combined.",
                "MediaPipe facial landmark detection was integrated.",
                "Mobile liveness processing strategy was defined."
            ]
        },
        {
            "num": 8, "date": "05.12.2025", "location": "Online (Microsoft Teams)",
            "period": "[UPDATE HOURS]", "attendees": "All team members, Advisor",
            "objectives": "UI/UX design review across platforms",
            "decisions": [
                "Web dashboard feature set was discussed.",
                "Mobile application screen flows were discussed.",
                "Desktop application modes were defined.",
                "UI/UX consistency across platforms was established."
            ]
        },
        {
            "num": 9, "date": "12.12.2025", "location": "Online (Microsoft Teams)",
            "period": "[UPDATE HOURS]", "attendees": "All team members, Advisor",
            "objectives": "Service integration and API contracts",
            "decisions": [
                "Identity Core and Biometric Processor integration strategy was completed.",
                "API contracts and standardized error responses were finalized.",
                "Integration and end-to-end testing plan was defined."
            ]
        },
        {
            "num": 10, "date": "26.12.2025", "location": "Online (Microsoft Teams)",
            "period": "[UPDATE HOURS]", "attendees": "All team members, Advisor",
            "objectives": "Documentation review and spring planning",
            "decisions": [
                "Documentation was reviewed and revised.",
                "Completed system components were demonstrated.",
                "Performance metrics were identified.",
                "Spring semester focus areas were clarified."
            ]
        },
        {
            "num": 11, "date": "09.01.2026", "location": "Marmara University, Engineering Faculty",
            "period": "[UPDATE HOURS]", "attendees": "All team members, Advisor",
            "objectives": "Fall semester retrospective and ADD review",
            "decisions": [
                "Fall semester retrospective was conducted.",
                "Spring semester goals and priorities were defined.",
                "Service integration tasks were prioritized.",
                "ADD document requirements were reviewed with the advisor.",
                "Mandatory diagrams and performance metrics were clarified."
            ]
        }
    ]

    for m in meetings:
        # Meeting header
        p = doc.add_paragraph()
        p.add_run(f"Meeting #{m['num']}").bold = True

        # Details table
        details = [
            ("Date:", m["date"]),
            ("Location:", m["location"]),
            ("Period:", m["period"]),
            ("Attendees:", m["attendees"]),
            ("Objectives:", m["objectives"])
        ]

        for label, value in details:
            p = doc.add_paragraph()
            p.add_run(f"    {label} ").bold = True
            p.add_run(value)

        # Decisions
        p = doc.add_paragraph()
        p.add_run("    Decisions and Notes:").bold = True
        for decision in m["decisions"]:
            doc.add_paragraph(f"        - {decision}")

        doc.add_paragraph()

    # ==========================================================
    # 4. TASK PLAN (Section 6.3) - Clean Tables + Image Gantt
    # ==========================================================
    doc.add_heading('4. Replacement for Section 6.3 (Task Plan)', level=1)
    doc.add_paragraph('Replace Section 6.3 with this version. Includes "Expected Output" column and professional Gantt chart images.')
    doc.add_paragraph()

    doc.add_heading('6.3 Task Plan', level=2)

    # Fall Semester
    doc.add_heading('6.3.1 Fall Semester Timeline', level=3)

    fall_data = [
        ['Task No.', 'Task Description', 'Expected Output', 'Sep', 'Oct', 'Nov', 'Dec', 'Jan'],
        ['F-1', 'Project Initialization & Setup', 'Git repository, CI/CD workflows, Literature summary', '', '', '', '', ''],
        ['F-2', 'Database Schema Design', 'ER diagram, Flyway migrations, pgvector config', '', '', '', '', ''],
        ['F-3', 'Identity Core - Base Implementation', 'Authentication API, JWT handling, User CRUD', '', '', '', '', ''],
        ['F-4', 'Biometric Processor - Core API', '46+ REST endpoints, Face detection module', '', '', '', '', ''],
        ['F-5', 'Liveness Detection Algorithm', 'Biometric Puzzle, EAR/MAR detection', '', '', '', '', ''],
        ['F-6', 'Web Admin Dashboard', 'React dashboard with 7 pages', '', '', '', '', ''],
        ['F-7', 'Service Integration', 'Inter-service communication', '', '', '', '', ''],
        ['F-8', 'Mobile App', 'Android app with camera access', '', '', '', '', ''],
        ['F-9', 'Desktop App', 'Kiosk mode prototype', '', '', '', '', ''],
        ['F-10', 'NFC Reader', 'NFC reading module', '', '', '', '', '']
    ]

    # Define which cells to shade (row, col) - 0-indexed, cols 3-7 are Sep-Jan
    fall_shading = {
        (1, 3): 'B4C6E7',  # F-1: Sep
        (2, 3): 'B4C6E7', (2, 4): 'B4C6E7',  # F-2: Sep-Oct
        (3, 4): 'B4C6E7', (3, 5): 'B4C6E7',  # F-3: Oct-Nov
        (4, 4): 'B4C6E7', (4, 5): 'B4C6E7',  # F-4: Oct-Nov
        (5, 5): 'B4C6E7', (5, 6): 'B4C6E7',  # F-5: Nov-Dec
        (6, 5): 'B4C6E7', (6, 6): 'B4C6E7',  # F-6: Nov-Dec
        (7, 6): 'B4C6E7', (7, 7): 'B4C6E7',  # F-7: Dec-Jan
        (8, 6): 'B4C6E7', (8, 7): 'B4C6E7',  # F-8: Dec-Jan
        (9, 7): 'B4C6E7',  # F-9: Jan
        (10, 7): 'B4C6E7',  # F-10: Jan
    }

    fall_table = doc.add_table(rows=len(fall_data), cols=8)
    fall_table.style = 'Table Grid'
    for i, row_data in enumerate(fall_data):
        row = fall_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(row.cells[j], 'D9E2F3')
            elif (i, j) in fall_shading:
                set_cell_shading(row.cells[j], fall_shading[(i, j)])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Table 18: Completed Tasks in Fall Semester').italic = True
    p = doc.add_paragraph()
    p.add_run('Note: ').bold = True
    p.add_run('Shaded cells indicate active development periods.')
    doc.add_paragraph()

    # Insert Fall Gantt Chart Image
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Figure: Fall Semester Gantt Chart').bold = True
    try:
        doc.add_picture('gantt_fall.png', width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('[Insert gantt_fall.png image here]')
    doc.add_paragraph()

    # Spring Semester
    doc.add_heading('6.3.2 Spring Semester Timeline', level=3)

    spring_data = [
        ['Task No.', 'Task Description', 'Expected Output', 'Feb', 'Mar', 'Apr', 'May', 'Jun'],
        ['S-1', 'RBAC Implementation', 'Complete RBAC system, Permission API', '', '', '', '', ''],
        ['S-2', 'Service Integration - Complete', 'Webhooks, Event-driven communication', '', '', '', '', ''],
        ['S-3', 'Mobile App - Production Ready', 'Play Store ready build', '', '', '', '', ''],
        ['S-4', 'Desktop App - Production Ready', 'Windows/Linux/macOS installers', '', '', '', '', ''],
        ['S-5', 'End-to-End Testing', 'Automated test suites', '', '', '', '', ''],
        ['S-6', 'Performance Optimization', 'Load test reports', '', '', '', '', ''],
        ['S-7', 'Security Audit', 'Security audit report', '', '', '', '', '']
    ]

    spring_shading = {
        (1, 3): 'FFF2CC',  # S-1: Feb
        (2, 3): 'FFF2CC', (2, 4): 'FFF2CC',  # S-2: Feb-Mar
        (3, 3): 'FFF2CC', (3, 4): 'FFF2CC',  # S-3: Feb-Mar
        (4, 4): 'FFF2CC', (4, 5): 'FFF2CC',  # S-4: Mar-Apr
        (5, 4): 'FFF2CC', (5, 5): 'FFF2CC',  # S-5: Mar-Apr
        (6, 5): 'FFF2CC', (6, 6): 'FFF2CC',  # S-6: Apr-May
        (7, 6): 'FFF2CC', (7, 7): 'FFF2CC',  # S-7: May-Jun
    }

    spring_table = doc.add_table(rows=len(spring_data), cols=8)
    spring_table.style = 'Table Grid'
    for i, row_data in enumerate(spring_data):
        row = spring_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(row.cells[j], 'D9E2F3')
            elif (i, j) in spring_shading:
                set_cell_shading(row.cells[j], spring_shading[(i, j)])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Table 19: Planned Tasks for Spring Semester').italic = True
    doc.add_paragraph()

    # Insert Spring Gantt Chart Image
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Figure: Spring Semester Gantt Chart').bold = True
    try:
        doc.add_picture('gantt_spring.png', width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        doc.add_paragraph('[Insert gantt_spring.png image here]')

    # ==========================================================
    # FOOTER
    # ==========================================================
    doc.add_paragraph()
    doc.add_paragraph('=' * 60)
    doc.add_paragraph('END OF FIXES DOCUMENT')
    doc.add_paragraph('Generated for FIVUCSAS ADD - CSE4197 Engineering Project')

    # Save
    doc.save('ADD_FIXES_FINAL.docx')
    print("\nSuccessfully created ADD_FIXES_FINAL.docx")
    print("\nFiles generated:")
    print("  - ADD_FIXES_FINAL.docx (Word document with all fixes)")
    print("  - gantt_fall.png (Fall semester Gantt chart)")
    print("  - gantt_spring.png (Spring semester Gantt chart)")
    print("\nRemember to:")
    print("  1. Update [UPDATE HOURS] placeholders with actual meeting durations")
    print("  2. Verify meeting locations are correct")

if __name__ == "__main__":
    create_improved_add_fixes()
